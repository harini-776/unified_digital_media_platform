"""Build TRUSTMEDIA_JOURNAL_PAPER.docx from the markdown source.

Recreates the SLM reference paper's layout conventions:
  - Title block (centred title + author + guide on the cover)
  - Numbered section headings ("1.", "1.1", "2.", etc.)
  - Embedded figures at "Figure. N – ..." anchors
  - Tables rendered as proper Word tables with header row + grid borders
  - Code blocks in monospace with light-grey shading
  - References as hanging-indent paragraphs

Markdown source : docs/TRUSTMEDIA_JOURNAL_PAPER.md
DOCX output     : docs/TRUSTMEDIA_JOURNAL_PAPER.docx
Figures         : docs/figures/*.png (referenced via ![alt](path))

Run:
    docs/figures/.figvenv/bin/python3 docs/build_docx.py
"""

from __future__ import annotations

import os
import re
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

DOCS = os.path.dirname(os.path.abspath(__file__))
SOURCE = os.path.join(DOCS, "TRUSTMEDIA_JOURNAL_PAPER.md")
OUT = os.path.join(DOCS, "TRUSTMEDIA_JOURNAL_PAPER.docx")

FIG_WIDTH = Inches(6.0)
BODY_FONT = "Times New Roman"
MONO_FONT = "Courier New"
TITLE_COLOR = RGBColor(0x1F, 0x2D, 0x3D)


def _force_font(run, name: str):
    """Set run font across all four OOXML slots (ascii / hAnsi / eastAsia / cs).

    `python-docx`'s run.font.name only updates the Latin slot, which leaves
    em-dashes, smart quotes, and CJK characters rendered in the document
    default. Setting all four slots ensures Word/LibreOffice respects the
    chosen font for every glyph.
    """
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{slot}"), name)


# ── Markdown parsing ────────────────────────────────────────────────────────

IMG_RE = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<src>[^)]+)\)\s*$")
HEADING_RE = re.compile(r"^(?P<n>\d+(?:\.\d+)?)\.?\s+(?P<text>.+)$")


def _read_blocks(path: str) -> list[tuple[str, list[str]]]:
    """Group the markdown into (kind, lines) blocks.

    kind ∈ { 'para', 'image', 'code', 'table', 'figure-anchor' }.
    Tables here are the SLM-style vertical-cell-stack: a sequence of single
    short lines separated by blank lines, NOT markdown grid tables.
    """
    with open(path) as f:
        raw = f.read().splitlines()

    blocks: list[tuple[str, list[str]]] = []
    buf: list[str] = []
    in_code = False

    def flush(kind: str = "para"):
        nonlocal buf
        if buf:
            blocks.append((kind, buf))
            buf = []

    for line in raw:
        if line.strip().startswith(">> CODE BLOCK <<"):
            if in_code:
                flush("code")
                in_code = False
            else:
                flush()
                in_code = True
            continue

        if in_code:
            buf.append(line)
            continue

        if not line.strip():
            flush()
            continue

        m = IMG_RE.match(line)
        if m:
            flush()
            blocks.append(("image", [m.group("src"), m.group("alt")]))
            continue

        buf.append(line)

    flush("code" if in_code else "para")
    return blocks


# ── Document plumbing ───────────────────────────────────────────────────────

def _set_cell_border(cell, **kwargs):
    """Apply borders to a docx table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            border = tc_borders.find(qn(f"w:{edge}"))
            if border is None:
                border = OxmlElement(f"w:{edge}")
                tc_borders.append(border)
            for k, v in kwargs[edge].items():
                border.set(qn(f"w:{k}"), v)


def _shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_bordered_table(doc: Document, rows: list[list[str]],
                        header: bool = True, widths: list[float] | None = None):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(cell_text)
            _force_font(run, BODY_FONT)
            run.font.size = Pt(10)
            if header and i == 0:
                run.bold = True
                _shade_cell(cell, "D9E2EC")
            border_spec = {"sz": "4", "val": "single", "color": "7B8794"}
            _set_cell_border(
                cell,
                top=border_spec, left=border_spec,
                bottom=border_spec, right=border_spec,
            )
            if widths and j < len(widths):
                cell.width = Inches(widths[j])

    if widths:
        for row in table.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)


# ── Inline emphasis (bold/italic/code) ──────────────────────────────────────

INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|"      # bold
    r"\*[^*]+\*|"            # italic
    r"`[^`]+`)"              # inline code
)


def _add_runs_with_inline(p, text: str, base_size: Pt = Pt(11)):
    pieces = INLINE_RE.split(text)
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            run = p.add_run(piece[2:-2])
            run.bold = True
            _force_font(run, BODY_FONT)
            run.font.size = base_size
        elif piece.startswith("*") and piece.endswith("*"):
            run = p.add_run(piece[1:-1])
            run.italic = True
            _force_font(run, BODY_FONT)
            run.font.size = base_size
        elif piece.startswith("`") and piece.endswith("`"):
            run = p.add_run(piece[1:-1])
            _force_font(run, MONO_FONT)
            run.font.size = Pt(10)
        else:
            run = p.add_run(piece)
            _force_font(run, BODY_FONT)
            run.font.size = base_size


# ── Front matter detection ──────────────────────────────────────────────────

# The reference's front matter is a sequence of short paragraphs before the
# Abstract: title, "By", author, register-number-style line, department,
# guide block, date. They are centred and styled differently from body.

FRONT_MATTER_BEFORE = "Abstract"

# Section headings: lines that match HEADING_RE AND are not part of front
# matter / not a known body line.
H1_PATTERN = re.compile(r"^(\d+)\.\s+(.+)$")
H2_PATTERN = re.compile(r"^(\d+)\.(\d+)\s+(.+)$")


# ── Tables: detect vertical-cell-stack pattern ──────────────────────────────
#
# The SLM reference renders tables as paragraphs of single short lines
# separated by blank lines. We can't auto-detect that reliably from the
# markdown text alone. Instead, the markdown source uses a literal table
# region delimited by short consecutive paragraphs. The simplest robust
# approach is: parse the document's table regions explicitly by their
# section/topic, since they are stable.
#
# Strategy: maintain a list of (section_title, table_rows) and emit those
# tables as bordered Word tables when the section is reached. Body prose
# inside the section continues to flow normally.


# Tables hard-coded by content (the source markdown uses the SLM-style
# vertical-stack representation, which doesn't round-trip cleanly through a
# generic markdown parser; we embed tabular data here so the .docx renders
# the proper tabular grid).

TABLES = {
    "3.2": {
        "after_marker": "3.2 Direct Comparison",
        "headers": ["Property", "Single-modality detector",
                    "Multimodal fusion (TrustMedia)"],
        "rows": [
            ["Independent signals", "1", "5"],
            ["Vulnerable to clean-modality attacks",
             "Yes (catastrophic)", "No (verdict needs all-modal flip)"],
            ["Cross-dataset AUC drop (FF++ → Celeb-DF)",
             "0.15–0.25 absolute", "0.05–0.10 absolute"],
            ["Coverage of audio-only manipulations",
             "None", "Full (voice + lipsync)"],
            ["Coverage of face-only manipulations", "Full", "Full"],
            ["Per-branch failure mode",
             "Total verdict failure", "Graceful degradation"],
            ["Calibrated probability output",
             "Rare", "Yes (temperature-scaled)"],
            ["Per-signal explainability", "Low", "Native (5 score cards)"],
            ["Robustness to missing modality",
             "None", "Yes (modality dropout p=0.2)"],
            ["Per-sample modality weights",
             "Not applicable", "Yes (learned attention)"],
            ["Provenance override available",
             "No", "Yes (on-chain SHA-256)"],
            ["Auditable failure mode",
             "End-to-end opaque", "Per-branch telemetry"],
            ["Suitable for forensic deployment", "Marginal", "Yes"],
        ],
        "widths": [2.4, 2.0, 2.6],
    },
    "2.2": {
        "after_marker": "2.2 Quantization-of-Effort: From Heuristics to Foundation Models",
        "headers": ["Tier", "Approach", "Per-branch effort", "Detection ceiling"],
        "rows": [
            ["Tier 0",
             "Hand-coded heuristics (Laplacian variance, EAR, Pearson correlation)",
             "Hours", "Modest, brittle"],
            ["Tier 1",
             "Tabular classifiers on engineered features (XGBoost on EAR, pose physics)",
             "Days", "Strong on biological signals"],
            ["Tier 2",
             "Trained-from-scratch CNNs / Transformers on FF++ scale corpora",
             "Weeks + GPU", "Strong in-distribution"],
            ["Tier 3",
             "Pre-trained foundation backbones (Wav2Vec2, ViT) fine-tuned on small corpora",
             "Days, no GPU required", "Strong cross-distribution"],
        ],
        "widths": [0.8, 3.4, 1.2, 1.6],
    },
    "6.1": {
        "after_marker": "6.1 Available Branches",
        "headers": ["Branch", "Architecture", "Backend selector", "Output"],
        "rows": [
            ["Face",
             "EfficientNet-B4 + Temporal Transformer (heuristic alt)",
             "FACE_BACKEND={heuristic, hf}",
             "face_score ∈ [0, 100]"],
            ["Voice",
             "Wav2Vec2 + MFCC CNN (heuristic alt)",
             "VOICE_BACKEND={heuristic, hf}",
             "voice_score ∈ [0, 100]"],
            ["Lipsync",
             "SyncNet-style dual-stream (heuristic alt)",
             "—",
             "lipsync_score ∈ [0, 100]"],
            ["Blink",
             "MediaPipe FaceMesh + EAR + XGBoost (n_estimators=200, depth=5)",
             "—",
             "blink_score ∈ [0, 100]"],
            ["HeadMotion",
             "solvePnP 3D pose + 18-D physics + XGBoost (n_estimators=200, depth=6)",
             "—",
             "headmotion_score ∈ [0, 100]"],
        ],
        "widths": [1.0, 2.6, 1.7, 1.7],
    },
    "6.2": {
        "after_marker": "6.2 Performance on Target Hardware",
        "headers": ["Branch", "Load time", "Per-30s-video latency", "Notes"],
        "rows": [
            ["Face (heuristic)", "<1s", "3–5s",
             "OpenCV + Laplacian + colour histogram"],
            ["Face (HF ViT)", "~6s", "15–25s",
             "prithivMLmods/Deep-Fake-Detector-v2-Model"],
            ["Voice (heuristic)", "<1s", "1–2s",
             "MFCC cross-segment correlation"],
            ["Voice (HF Wav2Vec2)", "~4s", "3–6s",
             "MelodyMachine/Deepfake-audio-detection-V2"],
            ["Lipsync (heuristic)", "<1s", "2–4s",
             "MediaPipe + Pearson correlation"],
            ["Blink", "<1s", "2–3s", "XGBoost on 14-D EAR features"],
            ["HeadMotion", "<1s", "3–5s",
             "XGBoost on 18-D physics features"],
        ],
        "widths": [1.6, 1.0, 1.6, 2.6],
    },
    "7.5": {
        "after_marker": "7.5 Uncertainty Quantification",
        "headers": ["Condition", "Flag", "Effect"],
        "rows": [
            ["H > 0.7 or Δ > 30", "HIGH",
             "Confidence × 0.75; flag in API response"],
            ["H > 0.3 or Δ > 15", "MEDIUM",
             "Confidence unchanged; flag in API response"],
            ["Otherwise", "LOW", "No effect"],
        ],
        "widths": [1.8, 1.0, 3.4],
    },
    "7.7": {
        "after_marker": "7.7 Verdict Mapping",
        "headers": ["p_calib", "Verdict", "Trust Score"],
        "rows": [
            ["≥ 0.70", "MANIPULATED", "0–30"],
            ["0.40 ≤ p < 0.70", "SUSPICIOUS", "30–60"],
            ["< 0.40", "AUTHENTIC", "60–100"],
        ],
        "widths": [1.6, 1.8, 1.6],
    },
    "9.1": {
        "after_marker": "9.1 API Contract",
        "headers": ["Method", "Path", "Description"],
        "rows": [
            ["POST", "/api/v1/auth/register",
             "Create user; returns access + refresh tokens"],
            ["POST", "/api/v1/auth/login",
             "Authenticate; returns JWT pair"],
            ["POST", "/api/v1/videos/upload",
             "Upload video; returns video_id, job_id"],
            ["GET", "/api/v1/jobs/{job_id}", "Poll analysis progress"],
            ["GET", "/api/v1/videos/{video_id}/result",
             "Retrieve full analysis result"],
            ["GET", "/api/v1/videos/{video_id}/stream",
             "Time-bounded signed-URL video stream"],
            ["POST", "/api/v1/blockchain/register",
             "Register media provenance on-chain"],
            ["POST", "/api/v1/blockchain/verify",
             "Verify hash against blockchain"],
        ],
        "widths": [0.9, 2.6, 3.0],
    },
    "12.1": {
        "after_marker": "12.1 Datasets",
        "headers": ["Dataset", "Real videos", "Manipulated videos", "Notes"],
        "rows": [
            ["FaceForensics++ (FF++) c23", "1,000", "4,000",
             "DeepFakes / Face2Face / FaceSwap / NeuralTextures, H.264 c23"],
            ["Celeb-DF v2", "590", "5,639",
             "Higher visual quality than FF++; standard cross-dataset target"],
            ["DFDC", "Diverse", "Diverse",
             "Large-scale, mixed manipulation methods and recording conditions"],
        ],
        "widths": [1.6, 0.9, 1.4, 2.7],
    },
    "12.2": {
        "after_marker": "12.2 Per-Branch Performance (FF++ c23, architecture)",
        "headers": ["Branch", "AUC-ROC", "EER (%)", "Notes"],
        "rows": [
            ["Face (EfficientNet-B4 + Temporal Transformer)",
             "0.91", "8.2", "Strongest single modality"],
            ["LipSync (SyncNet-style dual-stream)",
             "0.87", "10.1", "Strong on face-swap"],
            ["Voice (Wav2Vec2 + MFCC CNN)",
             "0.84", "12.3", "On audio-available subset"],
            ["Blink (EAR + XGBoost)",
             "0.79", "15.6", "Stronger on GAN-generated faces"],
            ["HeadMotion (solvePnP + XGBoost)",
             "0.82", "13.4", "Stronger on reenactment methods"],
            ["Fused (TrustMedia)",
             "0.94", "5.8", "All modalities"],
        ],
        "widths": [2.6, 0.9, 0.9, 2.0],
    },
    "12.3": {
        "after_marker": "12.3 Cross-Dataset Generalisation (Train: FF++, Test: Celeb-DF v2)",
        "headers": ["System", "AUC-ROC", "Drop vs. in-dist"],
        "rows": [
            ["Face-only (EfficientNet-B4)", "0.72", "−0.19"],
            ["LipSync-only", "0.68", "−0.19"],
            ["Voice-only", "0.77", "−0.07"],
            ["Static weighted fusion", "0.81", "−0.13"],
            ["TrustMedia (attention fusion)", "0.86", "−0.08"],
        ],
        "widths": [3.0, 1.5, 1.5],
    },
    "12.4": {
        "after_marker": "12.4 Calibration Quality",
        "headers": ["Model", "ECE pre-calibration",
                    "ECE post temperature scaling", "Temperature T"],
        "rows": [
            ["Face Branch", "0.089", "0.031", "1.24"],
            ["Fusion Model", "0.072", "0.021", "1.18"],
        ],
        "widths": [1.6, 1.7, 2.1, 1.2],
    },
    "12.5": {
        "after_marker": "12.5 Modality Ablation (FF++ c23)",
        "headers": ["Removed Modality", "AUC-ROC", "ΔAUC vs. Full"],
        "rows": [
            ["None (Full TrustMedia)", "0.941", "—"],
            ["− Face", "0.873", "−0.068"],
            ["− LipSync", "0.901", "−0.040"],
            ["− Voice", "0.908", "−0.033"],
            ["− Blink", "0.921", "−0.020"],
            ["− HeadMotion", "0.923", "−0.018"],
        ],
        "widths": [2.4, 1.4, 1.6],
    },
    "12.6": {
        "after_marker": "12.6 Blockchain Provenance Evaluation",
        "headers": ["Scenario", "N", "Correct verdict", "Notes"],
        "rows": [
            ["Registered authentic", "50", "50 (100%)",
             "On-chain hash matches; verdict overridden to AUTHENTIC"],
            ["Registered-but-tampered", "50", "50 (100%)",
             "On-chain hash mismatches uploaded hash; overridden to MANIPULATED"],
            ["Unregistered", "100", "94 (94%)",
             "AI pipeline verdict applies (AUC = 0.94)"],
        ],
        "widths": [1.6, 0.6, 1.2, 3.4],
    },
}


# ── Lines that the table data above replaces (so we don't double-render) ────
#
# Each table region in the source is a sequence of consecutive non-blank
# lines starting at the table's "after_marker" heading. We skip every line
# from the marker until either (a) the next section heading or (b) a blank-
# delimited prose paragraph that is clearly NOT part of the vertical cells.
# We approximate this by counting expected vertical-cell lines:
# n_cells = (1 + len(rows)) * n_columns. After we've consumed that many
# non-blank lines from the table region, normal flow resumes.

def _table_consumes_lines(spec: dict) -> int:
    """Total non-blank lines that make up a vertical-cell-stack table."""
    return (1 + len(spec["rows"])) * len(spec["headers"])


# ── Document builder ────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page setup: 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Default style: Times New Roman 11
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    # Force the East-Asian / hAnsi / cs font slots on the Normal style too,
    # so em-dashes and unicode punctuation in body text render in TNR.
    _normal_rPr = style.element.get_or_add_rPr()
    _normal_rFonts = _normal_rPr.find(qn("w:rFonts"))
    if _normal_rFonts is None:
        _normal_rFonts = OxmlElement("w:rFonts")
        _normal_rPr.append(_normal_rFonts)
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        _normal_rFonts.set(qn(f"w:{slot}"), BODY_FONT)

    # Heading styles - H1 = 24pt, H2 = 16pt, H3 = 14pt
    for level, size, bold in [(1, 24, True), (2, 16, True), (3, 14, True)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = BODY_FONT
        h.font.size = Pt(size)
        h.font.bold = bold
        h.font.color.rgb = TITLE_COLOR
        # Apply rFonts override on the heading style as well.
        _h_rPr = h.element.get_or_add_rPr()
        _h_rFonts = _h_rPr.find(qn("w:rFonts"))
        if _h_rFonts is None:
            _h_rFonts = OxmlElement("w:rFonts")
            _h_rPr.append(_h_rFonts)
        for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
            _h_rFonts.set(qn(f"w:{slot}"), BODY_FONT)

    with open(SOURCE) as f:
        lines = f.readlines()

    i = 0
    n = len(lines)

    # ── Front matter ────────────────────────────────────────────────────────
    # Lines until "Abstract" are the cover block. Render centred.

    front: list[str] = []
    while i < n and not lines[i].strip().startswith("Abstract"):
        s = lines[i].rstrip()
        if s:
            front.append(s)
        i += 1

    # First non-empty line of front is the title.
    if front:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(front[0])
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = TITLE_COLOR
        _force_font(run, BODY_FONT)

        # Remaining front-matter lines, centred, smaller font.
        for line in front[1:]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(12)
            _force_font(run, BODY_FONT)
            if line.strip().startswith("By") or line.startswith("Under the Guidance"):
                run.bold = True

    # ── Body ────────────────────────────────────────────────────────────────

    # Track which tables we've emitted so we don't re-emit and don't render
    # the source's vertical-cell-stack lines as paragraphs.
    table_skip_remaining = 0
    references_started = False

    while i < n:
        line = lines[i].rstrip()
        stripped = line.strip()

        # Blank line → spacer
        if not stripped:
            i += 1
            continue

        # While we are inside a vertical-cell-stack region replaced by a
        # rendered Word table, swallow non-blank lines.
        if table_skip_remaining > 0:
            table_skip_remaining -= 1
            i += 1
            continue

        # Image
        m_img = IMG_RE.match(stripped)
        if m_img:
            src = m_img.group("src")
            full = src if os.path.isabs(src) else os.path.join(DOCS, src)
            if os.path.exists(full):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(full, width=FIG_WIDTH)
            i += 1
            continue

        # Code block sentinel
        if stripped.startswith(">> CODE BLOCK <<"):
            i += 1
            code_lines: list[str] = []
            while i < n and not lines[i].strip().startswith(">> CODE BLOCK <<"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # skip closing sentinel
            for cl in code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(cl if cl else " ")
                _force_font(run, MONO_FONT)
                run.font.size = Pt(9.5)
            continue

        # Section heading - "N." or "N.N" prefix at start of line
        m_h2 = H2_PATTERN.match(stripped)
        m_h1 = H1_PATTERN.match(stripped)
        if m_h2:
            heading = doc.add_heading(stripped, level=2)
            run0 = heading.runs[0]
            run0.font.color.rgb = TITLE_COLOR
            run0.font.size = Pt(16)
            _force_font(run0, BODY_FONT)
            # If a known table follows this heading, emit it now and arrange
            # to skip the corresponding source lines.
            for spec in TABLES.values():
                if stripped.startswith(spec["after_marker"]):
                    rows = [spec["headers"]] + spec["rows"]
                    _add_bordered_table(doc, rows, header=True,
                                        widths=spec["widths"])
                    table_skip_remaining = _table_consumes_lines(spec)
                    break
            i += 1
            continue
        if m_h1:
            number = int(m_h1.group(1))
            text = m_h1.group(2)
            heading = doc.add_heading(stripped, level=1)
            run0 = heading.runs[0]
            run0.font.color.rgb = TITLE_COLOR
            run0.font.size = Pt(24)
            _force_font(run0, BODY_FONT)
            if text.strip().lower().startswith("references"):
                references_started = True
            i += 1
            continue

        # "References" plain heading (no number prefix)
        if stripped == "References":
            heading = doc.add_heading("References", level=1)
            run0 = heading.runs[0]
            run0.font.color.rgb = TITLE_COLOR
            run0.font.size = Pt(24)
            _force_font(run0, BODY_FONT)
            references_started = True
            i += 1
            continue

        # Abstract / Keywords specially styled as bold lead paragraph
        if stripped == "Abstract":
            h = doc.add_paragraph()
            run = h.add_run("Abstract")
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = TITLE_COLOR
            _force_font(run, BODY_FONT)
            i += 1
            continue
        if stripped == "Keywords:" or stripped.startswith("Keywords:"):
            p = doc.add_paragraph()
            run = p.add_run("Keywords: ")
            run.bold = True
            _force_font(run, BODY_FONT)
            rest = stripped[len("Keywords:"):].strip()
            run2 = p.add_run(rest)
            run2.italic = True
            _force_font(run2, BODY_FONT)
            i += 1
            continue

        # Figure caption ("Figure. N – ...")
        if stripped.startswith("Figure."):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.italic = True
            run.font.size = Pt(11)
            _force_font(run, BODY_FONT)
            i += 1
            continue

        # Reference entry inside the References section: hanging indent
        if references_started:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.3)
            _add_runs_with_inline(p, stripped)
            i += 1
            continue

        # 1.1 Contributions list - the SLM reference renders this as a
        # sequence of paragraphs without leading bullets. We do the same.
        # All other body paragraphs follow the same rule.

        # Default body paragraph
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(6)
        _add_runs_with_inline(p, stripped)
        i += 1

    # ── Save ────────────────────────────────────────────────────────────────
    doc.save(OUT)
    return OUT


def main():
    out = build()
    print(f"Wrote {out}")
    print(f"Size: {os.path.getsize(out):,} bytes")


if __name__ == "__main__":
    main()
